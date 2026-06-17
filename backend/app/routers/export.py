from __future__ import annotations

import re
from io import BytesIO

from fastapi import APIRouter
from fastapi.responses import StreamingResponse

from ..schemas.models import ExportFormat, ExportRequest

router = APIRouter(tags=["Export"])


def _safe_filename(stem: str) -> str:
	cleaned = "".join(ch for ch in stem.strip() if ch.isalnum() or ch in {" ", "-", "_"})
	return "_".join(cleaned.split()) or "study_aid"


def _strip_markdown(value: str) -> str:
	cleaned = value.strip()
	cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
	cleaned = re.sub(r"__(.*?)__", r"\1", cleaned)
	cleaned = cleaned.replace("`", "")
	cleaned = cleaned.strip(" *_")
	return cleaned.strip()


def _split_table_row(line: str) -> list[str]:
	return [_strip_markdown(cell) for cell in line.strip().strip("|").split("|")]


def _markdown_to_pdf_text(content: str) -> str:
	output: list[str] = []
	for raw_line in (content or "").splitlines():
		line = raw_line.strip()
		if not line:
			if output and output[-1] != "":
				output.append("")
			continue

		if line.startswith("|") and line.endswith("|"):
			cells = _split_table_row(line)
			is_separator = all(set(cell.replace(" ", "")) <= {"-", ":"} for cell in cells)
			is_header = [cell.lower() for cell in cells[:4]] == [
				"term",
				"type",
				"definition",
				"why it matters",
			]
			if is_separator or is_header:
				continue
			if len(cells) >= 4:
				term, term_type, definition, why_it_matters = cells[:4]
				if output and output[-1] != "":
					output.append("")
				label = term if not term_type else f"{term} ({term_type})"
				output.extend(
					[
						label,
						f"Definition: {definition}",
						f"Why it matters: {why_it_matters}",
					]
				)
				continue

		if line.startswith("#"):
			output.append(line.lstrip("#").strip())
			continue

		if line.startswith(('-', '*')):
			output.append(f"- {_strip_markdown(line[1:])}")
			continue

		output.append(_strip_markdown(line))

	return "\n".join(output).strip()


def _pdf_safe_text(value: str) -> str:
	replacements = {
		"\u2018": "'",
		"\u2019": "'",
		"\u201c": '"',
		"\u201d": '"',
		"\u2013": "-",
		"\u2014": "-",
		"\u2022": "-",
		"\u2026": "...",
		"\u00a0": " ",
	}
	cleaned = value or ""
	for source, replacement in replacements.items():
		cleaned = cleaned.replace(source, replacement)
	return cleaned.encode("latin-1", errors="replace").decode("latin-1")


def _build_pdf_bytes(title: str, content: str) -> bytes:
	from fpdf import FPDF

	pdf = FPDF()
	pdf.set_auto_page_break(auto=True, margin=15)
	pdf.add_page()

	effective_width = getattr(pdf, "epw", pdf.w - pdf.l_margin - pdf.r_margin)

	pdf.set_font("Helvetica", size=14)
	pdf.multi_cell(effective_width, 10, _pdf_safe_text(title))

	pdf.ln(2)
	pdf.set_font("Helvetica", size=11)
	pdf.multi_cell(effective_width, 6, _pdf_safe_text(_markdown_to_pdf_text(content)))

	out = pdf.output(dest="S")
	if isinstance(out, str):
		return out.encode("latin-1")
	return bytes(out)


def _build_docx_bytes(title: str, content: str) -> bytes:
	from docx import Document

	doc = Document()
	doc.add_heading(title, level=1)
	doc.add_paragraph(content)

	buf = BytesIO()
	doc.save(buf)
	return buf.getvalue()


@router.post(
	"/export",
	summary="Export a study aid as PDF or Word",
)
async def export(payload: ExportRequest) -> StreamingResponse:
	"""Exports a study aid as a downloadable PDF or DOCX."""

	title = payload.title or "Study Aid"
	safe_stem = _safe_filename(title)

	if payload.format == ExportFormat.pdf:
		data = _build_pdf_bytes(title=title, content=payload.content)
		filename = f"{safe_stem}.pdf"
		media_type = "application/pdf"
	else:
		data = _build_docx_bytes(title=title, content=payload.content)
		filename = f"{safe_stem}.docx"
		media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

	headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
	return StreamingResponse(BytesIO(data), media_type=media_type, headers=headers)
